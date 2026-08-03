"""Document text extraction.

The brief explicitly waives production-grade OCR, so this is deliberately
simple: pull the text layer out of whatever the user uploads and hand it to the
LLM. The intelligence lives in the extraction prompt, not here.

Supported: .pdf (pypdf), .docx (python-docx), .txt / .eml / .msg-as-text, and
raw pasted text through the same code path.
"""

import io
import logging
from email import policy
from email.parser import BytesParser

logging.getLogger("pypdf").setLevel(logging.ERROR)
logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".eml", ".md"}


class UnsupportedDocument(Exception):
    pass


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:  # noqa: BLE001
            logger.warning("Page %d of PDF failed to extract: %s", i, exc)
    return "\n\n".join(p for p in pages if p.strip())


def _from_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    blocks = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def _from_eml(data: bytes) -> str:
    msg = BytesParser(policy=policy.default).parsebytes(data)
    header = "\n".join(
        f"{k}: {msg[k]}" for k in ("From", "To", "Subject", "Date") if msg[k]
    )
    body = msg.get_body(preferencelist=("plain", "html"))
    content = body.get_content() if body else ""
    return f"{header}\n\n{content}"


def extract_text(filename: str, data: bytes) -> str:
    """Dispatch on file extension and return plain text."""
    name = (filename or "").lower()

    try:
        if name.endswith(".pdf"):
            text = _from_pdf(data)
        elif name.endswith(".docx"):
            text = _from_docx(data)
        elif name.endswith(".eml"):
            text = _from_eml(data)
        elif name.endswith((".txt", ".md")) or not name:
            text = data.decode("utf-8", errors="replace")
        else:
            raise UnsupportedDocument(
                f"Unsupported file type. Accepted formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )
    except UnsupportedDocument:
        raise
    except Exception as exc:  # noqa: BLE001
        logger.exception("Extraction failed for %s", filename)
        raise UnsupportedDocument(f"Could not read {filename}: {exc}") from exc

    cleaned = "\n".join(line.rstrip() for line in text.splitlines())
    while "\n\n\n" in cleaned:
        cleaned = cleaned.replace("\n\n\n", "\n\n")
    return cleaned.strip()
